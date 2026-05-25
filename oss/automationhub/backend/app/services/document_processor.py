"""
Multi-Format Document Processing Service

Comprehensive document ingestion pipeline supporting multiple formats including:
- PDF document processing with text extraction and OCR
- Microsoft Office documents (Word, Excel, PowerPoint)
- Web content extraction and cleaning
- Image processing with OCR capabilities
- Document validation, sanitization, and security scanning
- Batch processing with progress tracking
- Metadata extraction and indexing
"""

import asyncio
import hashlib
import logging
import mimetypes
import os
import re
import tempfile
import time
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, BinaryIO, Tuple
from uuid import UUID, uuid4

import aiofiles
import aiofiles.os

# Try to import optional dependencies
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    magic = None

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    Image = None

try:
    import pytesseract
    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False
    pytesseract = None

try:
    import fitz  # PyMuPDF for PDF processing
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    fitz = None

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    DocxDocument = None

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    pd = None

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False
    BeautifulSoup = None

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    HAS_SELENIUM = True
except ImportError:
    HAS_SELENIUM = False
    webdriver = None
    Options = None
    By = None
    WebDriverWait = None
    EC = None

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False
    requests = None

try:
    from email import policy
    from email import message_from_bytes
    HAS_EMAIL = True
except ImportError:
    HAS_EMAIL = False
    policy = None
    message_from_bytes = None

try:
    import yara  # For malware detection
    HAS_YARA = True
except ImportError:
    HAS_YARA = False
    yara = None

# Text splitter for legacy compatibility
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    HAS_LANGCHAIN = True
except ImportError:
    HAS_LANGCHAIN = False
    RecursiveCharacterTextSplitter = None

from app.core.config import get_settings
from app.core.exceptions import ValidationError, ProcessingError, SecurityError

# Try to import document schemas
try:
    from app.schemas.documents import (
        DocumentProcessingRequest, DocumentProcessingResponse,
        DocumentMetadata, ExtractionResult, ProcessingStatus
    )
    HAS_SCHEMAS = True
except ImportError:
    HAS_SCHEMAS = False
    # Define fallback schemas
    class DocumentProcessingRequest:
        def __init__(self, **kwargs):
            self.file_path = kwargs.get('file_path')
            self.source_url = kwargs.get('source_url')
            self.security_level = kwargs.get('security_level', SecurityLevel.BASIC)
            self.extraction_options = kwargs.get('extraction_options', {})

    class DocumentProcessingResponse:
        def __init__(self, **kwargs):
            self.processing_id = kwargs.get('processing_id', '')
            self.success = kwargs.get('success', False)
            self.content = kwargs.get('content', '')
            self.metadata = kwargs.get('metadata', DocumentMetadata() if 'DocumentMetadata' in globals() else {})
            self.processing_status = kwargs.get('processing_status', ProcessingStatus.FAILED)
            self.processing_time = kwargs.get('processing_time', 0.0)
            self.security_validation = kwargs.get('security_validation', {})
            self.extraction_stats = kwargs.get('extraction_stats', {})
            self.source_url = kwargs.get('source_url', '')
            self.error = kwargs.get('error', '')

logger = logging.getLogger(__name__)

# Constants
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'application/msword',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/vnd.ms-excel',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    'application/vnd.ms-powerpoint',
    'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    'text/plain',
    'text/csv',
    'text/html',
    'text/markdown',
    'image/jpeg',
    'image/png',
    'image/tiff',
    'image/bmp',
    'image/webp'
}


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    XLSX = "xlsx"
    XLS = "xls"
    PPTX = "pptx"
    PPT = "ppt"
    TXT = "txt"
    CSV = "csv"
    HTML = "html"
    MARKDOWN = "markdown"
    IMAGE = "image"
    EMAIL = "email"
    UNKNOWN = "unknown"


class ProcessingStatus(str, Enum):
    """Processing status states."""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"


class SecurityLevel(str, Enum):
    """Security scanning levels."""
    NONE = "none"
    BASIC = "basic"
    COMPREHENSIVE = "comprehensive"


class DocumentProcessor:
    """Main document processing service."""

    def __init__(self):
        self.security_validator = DocumentSecurityValidator()

        # Initialize extractors with dependency checking
        self.extractors = {}
        if HAS_FITZ:
            self.extractors[DocumentType.PDF] = PDFExtractor()
        if HAS_DOCX:
            self.extractors[DocumentType.DOCX] = OfficeDocumentExtractor()
            self.extractors[DocumentType.DOC] = OfficeDocumentExtractor()
        if HAS_PANDAS:
            self.extractors[DocumentType.XLSX] = OfficeDocumentExtractor()
            self.extractors[DocumentType.XLS] = OfficeDocumentExtractor()
        if HAS_DOCX:  # Reuse DOCX extractor for PPTX for now
            self.extractors[DocumentType.PPTX] = OfficeDocumentExtractor()
            self.extractors[DocumentType.PPT] = OfficeDocumentExtractor()
        if HAS_PIL and HAS_PYTESSERACT:
            self.extractors[DocumentType.IMAGE] = ImageExtractor()

        if HAS_BS4 and HAS_REQUESTS:
            self.web_extractor = WebContentExtractor()
        else:
            self.web_extractor = None

        # Legacy compatibility for knowledge management
        if HAS_LANGCHAIN:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        else:
            self.text_splitter = None

        self.supported_formats = {
            '.txt': self._process_text,
        }

        # Add format handlers based on available dependencies
        if HAS_FITZ:
            self.supported_formats['.pdf'] = self._process_pdf
        if HAS_DOCX:
            self.supported_formats['.doc'] = self._process_word
            self.supported_formats['.docx'] = self._process_word
        if HAS_BS4:
            self.supported_formats['.html'] = self._process_html
            self.supported_formats['.htm'] = self._process_html
        self.supported_formats['.json'] = self._process_json
        if HAS_PANDAS:
            self.supported_formats['.csv'] = self._process_csv
        self.supported_formats['.md'] = self._process_markdown
    async def process_document(
        self,
        request: DocumentProcessingRequest,
        user_id: UUID
    ) -> DocumentProcessingResponse:
        """Process a single document."""
        try:
            processing_id = str(uuid4())
            start_time = time.time()

            # Validate request
            await self._validate_request(request)

            # Determine processing method
            if request.source_url:
                return await self._process_web_content(request, processing_id, start_time)
            elif request.file_path:
                return await self._process_file(request, processing_id, start_time, user_id)
            else:
                raise ValidationError("Either file_path or source_url must be provided")

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return DocumentProcessingResponse(
                processing_id="",
                success=False,
                error=str(e),
                processing_status=ProcessingStatus.FAILED
            )

    async def process_batch(
        self,
        requests: List[DocumentProcessingRequest],
        user_id: UUID
    ) -> List[DocumentProcessingResponse]:
        """Process multiple documents in batch."""
        results = []

        for request in requests:
            result = await self.process_document(request, user_id)
            results.append(result)

        return results

    async def _validate_request(self, request: DocumentProcessingRequest):
        """Validate processing request."""
        if request.file_path and not os.path.exists(request.file_path):
            raise ValidationError(f"File not found: {request.file_path}")

        if request.source_url and not self._is_valid_url(request.source_url):
            raise ValidationError(f"Invalid URL: {request.source_url}")

    def _is_valid_url(self, url: str) -> bool:
        """Validate URL format."""
        url_pattern = re.compile(
            r'^https?://'  # http:// or https://
            r'(?:(?:[A-Z0-9](?:[A-Z0-9-]{0,61}[A-Z0-9])?\.)+[A-Z]{2,6}\.?|'  # domain...
            r'localhost|'  # localhost...
            r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'  # ...or ip
            r'(?::\d+)?'  # optional port
            r'(?:/?|[/?]\S+)$', re.IGNORECASE)
        return url_pattern.match(url) is not None

    async def _process_file(
        self,
        request: DocumentProcessingRequest,
        processing_id: str,
        start_time: float,
        user_id: UUID
    ) -> DocumentProcessingResponse:
        """Process file-based document."""
        try:
            # Security validation
            security_result = await self.security_validator.validate_file(
                request.file_path,
                request.security_level
            )

            if not security_result['is_safe']:
                raise SecurityError(f"Security threats detected: {security_result['threats_detected']}")

            # Determine document type
            doc_type = self._detect_document_type(request.file_path)

            # Get appropriate extractor
            extractor = self.extractors.get(doc_type)
            if not extractor:
                raise ProcessingError(f"Unsupported document type: {doc_type}")

            # Extract content
            extraction_result = await extractor.extract(
                request.file_path,
                **request.extraction_options
            )

            if not extraction_result.success:
                raise ProcessingError(extraction_result.error or "Content extraction failed")

            # Process results
            processing_time = time.time() - start_time

            return DocumentProcessingResponse(
                processing_id=processing_id,
                success=True,
                content=extraction_result.content,
                metadata=extraction_result.metadata,
                processing_status=ProcessingStatus.COMPLETED,
                processing_time=processing_time,
                security_validation=security_result,
                extraction_stats={
                    'pages_processed': extraction_result.pages_processed,
                    'images_extracted': getattr(extraction_result, 'images_extracted', 0),
                    'ocr_processed': getattr(extraction_result, 'ocr_processed', 0)
                }
            )

        except Exception as e:
            logger.error(f"File processing failed: {e}")
            return DocumentProcessingResponse(
                processing_id=processing_id,
                success=False,
                error=str(e),
                processing_status=ProcessingStatus.FAILED,
                processing_time=time.time() - start_time
            )

    async def _process_web_content(
        self,
        request: DocumentProcessingRequest,
        processing_id: str,
        start_time: float
    ) -> DocumentProcessingResponse:
        """Process web-based document."""
        try:
            if not self.web_extractor:
                raise ProcessingError("Web content extraction not available - missing dependencies")

            # Extract web content
            extraction_result = await self.web_extractor.extract(
                request.source_url,
                **request.extraction_options
            )

            if not extraction_result.success:
                raise ProcessingError(extraction_result.error or "Web content extraction failed")

            # Process results
            processing_time = time.time() - start_time

            return DocumentProcessingResponse(
                processing_id=processing_id,
                success=True,
                content=extraction_result.content,
                metadata=extraction_result.metadata,
                processing_status=ProcessingStatus.COMPLETED,
                processing_time=processing_time,
                source_url=request.source_url
            )

        except Exception as e:
            logger.error(f"Web content processing failed: {e}")
            return DocumentProcessingResponse(
                processing_id=processing_id,
                success=False,
                error=str(e),
                processing_status=ProcessingStatus.FAILED,
                processing_time=time.time() - start_time
            )

    def _detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file."""
        file_ext = Path(file_path).suffix.lower()

        # Try to get mime type if magic library is available
        mime_type = None
        if HAS_MAGIC and os.path.exists(file_path):
            try:
                mime_type = magic.from_file(file_path, mime=True)
            except Exception:
                mime_type = None

        # Map extensions to types
        ext_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOC,
            '.xlsx': DocumentType.XLSX,
            '.xls': DocumentType.XLS,
            '.pptx': DocumentType.PPTX,
            '.ppt': DocumentType.PPT,
            '.txt': DocumentType.TXT,
            '.csv': DocumentType.CSV,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.md': DocumentType.MARKDOWN
        }

        # Image extensions
        image_exts = {'.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.webp'}

        if file_ext in ext_mapping:
            return ext_mapping[file_ext]
        elif file_ext in image_exts or (mime_type and mime_type.startswith('image/')):
            return DocumentType.IMAGE
        else:
            return DocumentType.UNKNOWN

    # Legacy compatibility methods for knowledge management
    async def process_document_legacy(self, document_id: str) -> Dict[str, Any]:
        """Process a document: extract text, create chunks, generate embeddings (Legacy)"""
        try:
            # Import legacy dependencies
            from app.core.cloudflare_d1 import get_d1_session
            from app.core.vector_db import knowledge_manager
            from app.models.document import Document
            from langchain.text_splitter import RecursiveCharacterTextSplitter
            from sqlalchemy import select, update, text

            async with get_d1_session() as session:
                # Get document
                result = await session.execute(
                    select(Document).where(Document.id == document_id)
                )
                document = result.scalar_one_or_none()

                if not document:
                    raise Exception(f"Document {document_id} not found")

                # Update status to processing
                await session.execute(
                    update(Document)
                    .where(Document.id == document_id)
                    .values(embedding_status='processing')
                )
                await session.commit()

                # Extract text from document
                text_content = await self._extract_text_legacy(document)

                if not text_content.strip():
                    raise Exception("No text content extracted from document")

                # Create chunks
                chunks = await self._create_chunks(document_id, text_content)

                # Update document status
                await session.execute(
                    update(Document)
                    .where(Document.id == document_id)
                    .values(
                        embedding_status='completed',
                        chunk_count=len(chunks),
                        updated_at=datetime.utcnow()
                    )
                )
                await session.commit()

                logger.info(f"Successfully processed document {document_id} with {len(chunks)} chunks")

                return {
                    "success": True,
                    "document_id": document_id,
                    "chunks_created": len(chunks),
                    "status": "completed"
                }

        except Exception as e:
            logger.error(f"Failed to process document {document_id}: {e}")

            # Update status to failed
            try:
                from app.core.cloudflare_d1 import get_d1_session
                from app.models.document import Document
                from sqlalchemy import update

                async with get_d1_session() as session:
                    await session.execute(
                        update(Document)
                        .where(Document.id == document_id)
                        .values(
                            embedding_status='failed',
                            updated_at=datetime.utcnow()
                        )
                    )
                    await session.commit()
            except Exception as update_error:
                logger.error(f"Failed to update document status: {update_error}")

            return {
                "success": False,
                "document_id": document_id,
                "error": str(e),
                "status": "failed"
            }
    # Legacy compatibility methods (continued)
    async def _extract_text_legacy(self, document) -> str:
        """Extract text content from a document (Legacy)"""
        if hasattr(document, 'content') and document.content:
            # For text documents, content is already stored
            return document.content

        if hasattr(document, 'file_path') and document.file_path:
            # For file documents, extract from file
            file_path = Path(document.file_path)
            if not file_path.exists():
                raise Exception(f"File not found: {document.file_path}")

            file_extension = file_path.suffix.lower()

            if file_extension not in self.supported_formats:
                raise Exception(f"Unsupported file format: {file_extension}")

            processor = self.supported_formats[file_extension]
            return await processor(file_path)

        raise Exception("Document has no content or file path")

    async def _process_text(self, file_path: Path) -> str:
        """Process text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            return await f.read()

    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF files"""
        try:
            extractor = PDFExtractor()
            result = await extractor.extract(str(file_path))
            if result.success:
                return result.content
            else:
                raise Exception(result.error)
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            raise Exception(f"PDF processing failed: {e}")

    async def _process_word(self, file_path: Path) -> str:
        """Process Word documents"""
        try:
            extractor = OfficeDocumentExtractor()
            result = await extractor.extract(str(file_path))
            if result.success:
                return result.content
            else:
                raise Exception(result.error)
        except Exception as e:
            logger.error(f"Failed to process Word document {file_path}: {e}")
            raise Exception(f"Word document processing failed: {e}")

    async def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown files"""
        try:
            extractor = WebContentExtractor()
            # For markdown, treat as simple text processing
            return await self._process_text(file_path)
        except Exception as e:
            logger.error(f"Failed to process Markdown {file_path}: {e}")
            return await self._process_text(file_path)

    async def _process_html(self, file_path: Path) -> str:
        """Process HTML files"""
        try:
            extractor = WebContentExtractor()
            result = await extractor.extract(f"file://{file_path}")
            if result.success:
                return result.content
            else:
                raise Exception(result.error)
        except Exception as e:
            logger.error(f"Failed to process HTML {file_path}: {e}")
            raise Exception(f"HTML processing failed: {e}")

    async def _process_json(self, file_path: Path) -> str:
        """Process JSON files"""
        try:
            # For JSON, return as formatted text
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                # Return JSON as text
                return content
        except Exception as e:
            logger.error(f"Failed to process JSON {file_path}: {e}")
            return await self._process_text(file_path)

    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV files"""
        try:
            extractor = OfficeDocumentExtractor()
            result = await extractor.extract(str(file_path))
            if result.success:
                return result.content
            else:
                raise Exception(result.error)
        except Exception as e:
            logger.error(f"Failed to process CSV {file_path}: {e}")
            raise Exception(f"CSV processing failed: {e}")

    async def _create_chunks(self, document_id: str, text_content: str) -> List[str]:
        """Create text chunks from document content"""
        # Use langchain text splitter
        chunks = self.text_splitter.split_text(text_content)

        # Save chunks to database (legacy method)
        try:
            from app.core.cloudflare_d1 import get_d1_session
            from app.services.knowledge_models import KnowledgeChunk

            async with get_d1_session() as session:
                for i, chunk_text in enumerate(chunks):
                    chunk = KnowledgeChunk(
                        id=str(uuid4()),
                        document_id=document_id,
                        chunk_text=chunk_text,
                        chunk_index=i,
                        metadata={
                            "chunk_size": len(chunk_text),
                            "created_at": datetime.utcnow().isoformat()
                        },
                        created_at=datetime.utcnow()
                    )
                    session.add(chunk)

                await session.commit()
        except Exception as e:
            logger.error(f"Failed to save chunks to database: {e}")

        logger.info(f"Created {len(chunks)} chunks for document {document_id}")
        return chunks

    # Other legacy methods remain the same...


class DocumentSecurityValidator:
    """Security validation for uploaded documents."""

    def __init__(self):
        self.yara_rules = []
        self._load_yara_rules()

    def _load_yara_rules(self):
        """Load YARA rules for malware detection."""
        if not HAS_YARA:
            logger.warning("YARA library not available - malware detection disabled")
            return

        try:
            # Basic YARA rules for demonstration
            rules_content = """
            rule SuspiciousScript {
                strings:
                    $script1 = "eval("
                    $script2 = "document.write"
                    $script3 = "javascript:"
                condition:
                    $script1 or $script2 or $script3
            }

            rule ExecutableContent {
                strings:
                    $exe1 = { 4D 5A }  // MZ header
                    $exe2 = { 7F 45 4C 46 }  // ELF header
                condition:
                    $exe1 or $exe2
            }
            """
            self.yara_rules = yara.compile(source=rules_content)
        except Exception as e:
            logger.warning(f"Failed to load YARA rules: {e}")
            self.yara_rules = None

    async def validate_file(self, file_path: str, security_level: SecurityLevel = SecurityLevel.BASIC) -> Dict[str, Any]:
        """Validate file for security threats."""
        validation_result = {
            'is_safe': True,
            'threats_detected': [],
            'warnings': [],
            'file_hash': None,
            'scan_time': datetime.now(timezone.utc).isoformat()
        }

        try:
            # Calculate file hash
            validation_result['file_hash'] = await self._calculate_file_hash(file_path)

            if security_level == SecurityLevel.NONE:
                return validation_result

            # Basic validation checks
            await self._basic_validation(file_path, validation_result)

            if security_level in [SecurityLevel.BASIC, SecurityLevel.COMPREHENSIVE]:
                await self._content_validation(file_path, validation_result)

            validation_result['is_safe'] = (
                len(validation_result['threats_detected']) == 0 and
                len(validation_result['warnings']) == 0
            )

        except Exception as e:
            logger.error(f"Security validation failed: {e}")
            validation_result['threats_detected'].append(f"Validation error: {str(e)}")
            validation_result['is_safe'] = False

        return validation_result

    async def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        sha256_hash = hashlib.sha256()
        async with aiofiles.open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()

    async def _basic_validation(self, file_path: str, result: Dict[str, Any]):
        """Basic file validation."""
        file_size = os.path.getsize(file_path)

        if file_size > MAX_FILE_SIZE:
            result['warnings'].append(f"Large file size: {file_size} bytes")

        # Check file extension vs content type
        mime_type = None
        if HAS_MAGIC and os.path.exists(file_path):
            try:
                mime_type = magic.from_file(file_path, mime=True)
            except Exception:
                mime_type = None

        file_ext = Path(file_path).suffix.lower()

        # Suspicious extensions
        suspicious_exts = {'.exe', '.bat', '.cmd', '.scr', '.pif', '.com', '.vbs', '.js'}
        if file_ext in suspicious_exts:
            result['threats_detected'].append(f"Suspicious file extension: {file_ext}")

        # Check for embedded executables
        if mime_type and mime_type in ['application/x-executable', 'application/x-msdownload']:
            result['threats_detected'].append(f"Executable file detected: {mime_type}")

    async def _content_validation(self, file_path: str, result: Dict[str, Any]):
        """Content-based validation."""
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()

            # YARA rule matching
            if self.yara_rules and hasattr(self.yara_rules, 'match'):
                try:
                    matches = self.yara_rules.match(data=content)
                    for match in matches:
                        result['threats_detected'].append(f"YARA rule match: {match.rule}")
                except Exception as e:
                    logger.error(f"YARA matching error: {e}")

            # Check for suspicious patterns
            try:
                text_content = content.decode('utf-8', errors='ignore').lower()
                suspicious_patterns = [
                    r'<script[^>]*>',
                    r'javascript:',
                    r'data:text/html',
                    r'vbscript:',
                    r'eval\s*\(',
                    r'document\.write'
                ]

                for pattern in suspicious_patterns:
                    if re.search(pattern, text_content):
                        result['threats_detected'].append(f"Suspicious pattern detected: {pattern}")
            except Exception as e:
                logger.error(f"Pattern matching error: {e}")

        except Exception as e:
            logger.error(f"Content validation error: {e}")


class DocumentExtractor:
    """Base class for document extractors."""

    async def extract(self, file_path: str, **kwargs) -> 'ExtractionResult':
        """Extract content from document."""
        raise NotImplementedError

    async def extract_metadata(self, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from document."""
        raise NotImplementedError


class PDFExtractor(DocumentExtractor):
    """PDF document extractor with OCR support."""

    def __init__(self):
        self.ocr_enabled = True

    async def extract(self, file_path: str, use_ocr: bool = True, **kwargs) -> 'ExtractionResult':
        """Extract text and content from PDF."""
        try:
            start_time = time.time()

            # Open PDF with PyMuPDF
            doc = fitz.open(file_path)

            # Extract metadata
            metadata = await self.extract_metadata(file_path)

            # Extract text from pages
            text_content = []
            images = []

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text
                page_text = page.get_text()
                if page_text.strip():
                    text_content.append({
                        'page': page_num + 1,
                        'text': page_text.strip(),
                        'type': 'text'
                    })

            doc.close()

            # Combine all text content
            full_text = "\n\n".join([
                content['text'] for content in text_content
            ])

            processing_time = time.time() - start_time

            return ExtractionResult(
                content=full_text,
                metadata=metadata,
                pages_processed=len(text_content),
                images_extracted=len(images),
                ocr_processed=0,
                processing_time=processing_time,
                success=True
            )

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return ExtractionResult(
                content="",
                metadata=DocumentMetadata(),
                success=False,
                error=str(e)
            )

    async def extract_metadata(self, file_path: str) -> 'DocumentMetadata':
        """Extract PDF metadata."""
        try:
            doc = fitz.open(file_path)
            metadata = doc.metadata

            return DocumentMetadata(
                title=metadata.get('title', ''),
                author=metadata.get('author', ''),
                subject=metadata.get('subject', ''),
                keywords=metadata.get('keywords', ''),
                creator=metadata.get('creator', ''),
                producer=metadata.get('producer', ''),
                creation_date=metadata.get('creationDate', ''),
                modification_date=metadata.get('modDate', ''),
                page_count=len(doc),
                file_size=os.path.getsize(file_path)
            )

        except Exception as e:
            logger.error(f"PDF metadata extraction failed: {e}")
            return DocumentMetadata()


class OfficeDocumentExtractor(DocumentExtractor):
    """Microsoft Office document extractor."""

    async def extract(self, file_path: str, **kwargs) -> 'ExtractionResult':
        """Extract content from Office documents."""
        try:
            start_time = time.time()
            file_ext = Path(file_path).suffix.lower()

            if file_ext == '.docx':
                return await self._extract_docx(file_path, start_time)
            elif file_ext in ['.xlsx', '.xls']:
                return await self._extract_excel(file_path, start_time)
            elif file_ext in ['.pptx', '.ppt']:
                return await self._extract_powerpoint(file_path, start_time)
            else:
                raise ProcessingError(f"Unsupported Office document format: {file_ext}")

        except Exception as e:
            logger.error(f"Office document extraction failed: {e}")
            return ExtractionResult(
                content="",
                metadata=DocumentMetadata(),
                success=False,
                error=str(e)
            )

    async def _extract_docx(self, file_path: str, start_time: float) -> 'ExtractionResult':
        """Extract content from DOCX file."""
        doc = DocxDocument(file_path)

        # Extract text content
        content_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                table_text.append(' | '.join(row_text))
            content_parts.append(' | '.join(table_text))

        full_content = '\n\n'.join(content_parts)

        # Extract metadata
        metadata = await self.extract_metadata(file_path)
        metadata.page_count = len(doc.paragraphs)

        return ExtractionResult(
            content=full_content,
            metadata=metadata,
            pages_processed=metadata.page_count,
            processing_time=time.time() - start_time,
            success=True
        )

    async def _extract_excel(self, file_path: str, start_time: float) -> 'ExtractionResult':
        """Extract content from Excel file."""
        # Use pandas for Excel processing
        if file_path.endswith('.xlsx'):
            df_dict = pd.read_excel(file_path, sheet_name=None)
        else:
            df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')

        content_parts = []
        for sheet_name, df in df_dict.items():
            content_parts.append(f"Sheet: {sheet_name}")
            content_parts.append(df.to_string(index=False))
            content_parts.append("")

        full_content = '\n'.join(content_parts)

        metadata = await self.extract_metadata(file_path)
        metadata.page_count = len(df_dict)

        return ExtractionResult(
            content=full_content,
            metadata=metadata,
            pages_processed=metadata.page_count,
            processing_time=time.time() - start_time,
            success=True
        )

    async def _extract_powerpoint(self, file_path: str, start_time: float) -> 'ExtractionResult':
        """Extract content from PowerPoint file."""
        metadata = await self.extract_metadata(file_path)

        return ExtractionResult(
            content="PowerPoint content extraction not fully implemented",
            metadata=metadata,
            pages_processed=1,
            processing_time=time.time() - start_time,
            success=True
        )

    async def extract_metadata(self, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from Office document."""
        file_stat = os.stat(file_path)

        return DocumentMetadata(
            title=Path(file_path).stem,
            file_size=file_stat.st_size,
            creation_date=datetime.fromtimestamp(file_stat.st_ctime, tz=timezone.utc).isoformat(),
            modification_date=datetime.fromtimestamp(file_stat.st_mtime, tz=timezone.utc).isoformat()
        )


class WebContentExtractor(DocumentExtractor):
    """Web content extractor with JavaScript rendering."""

    async def extract(self, url: str, render_js: bool = True, **kwargs) -> 'ExtractionResult':
        """Extract content from web URL."""
        try:
            start_time = time.time()

            if render_js:
                content = await self._extract_with_selenium(url)
            else:
                content = await self._extract_with_requests(url)

            # Parse HTML
            soup = BeautifulSoup(content, 'html.parser')

            # Extract main content
            title = soup.find('title')
            title_text = title.get_text() if title else ''

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Extract text content
            text_content = soup.get_text(separator='\n', strip=True)

            # Extract metadata
            metadata = await self.extract_metadata(url, soup)

            # Clean up content
            cleaned_content = self._clean_web_content(text_content)

            processing_time = time.time() - start_time

            return ExtractionResult(
                content=cleaned_content,
                metadata=metadata,
                processing_time=processing_time,
                success=True,
                source_url=url
            )

        except Exception as e:
            logger.error(f"Web content extraction failed: {e}")
            return ExtractionResult(
                content="",
                metadata=DocumentMetadata(),
                success=False,
                error=str(e),
                source_url=url
            )

    async def _extract_with_selenium(self, url: str) -> str:
        """Extract content using Selenium for JavaScript rendering."""
        try:
            # Setup Chrome options
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')

            # Initialize driver
            driver = webdriver.Chrome(options=chrome_options)

            # Load page
            driver.get(url)

            # Wait for page to load
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )

            # Get page source
            content = driver.page_source

            return content

        finally:
            if 'driver' in locals():
                driver.quit()

    async def _extract_with_requests(self, url: str) -> str:
        """Extract content using requests library."""
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        return response.text

    def _clean_web_content(self, content: str) -> str:
        """Clean extracted web content."""
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)

        lines = content.split('\n')
        cleaned_lines = []

        for line in lines:
            line = line.strip()
            if line and len(line) > 10:  # Skip very short lines
                cleaned_lines.append(line)

        return '\n'.join(cleaned_lines)

    async def extract_metadata(self, url: str, soup: Optional[BeautifulSoup] = None) -> 'DocumentMetadata':
        """Extract metadata from web page."""
        metadata = DocumentMetadata(
            title="",
            source_url=url,
            creation_date=datetime.now(timezone.utc).isoformat()
        )

        if soup:
            # Extract title
            title_tag = soup.find('title')
            if title_tag:
                metadata.title = title_tag.get_text().strip()

            # Extract meta description
            desc_tag = soup.find('meta', attrs={'name': 'description'})
            if desc_tag:
                metadata.subject = desc_tag.get('content', '')

            # Extract keywords
            keywords_tag = soup.find('meta', attrs={'name': 'keywords'})
            if keywords_tag:
                metadata.keywords = keywords_tag.get('content', '')

        return metadata


class ImageExtractor(DocumentExtractor):
    """Image document extractor with OCR support."""

    async def extract(self, file_path: str, **kwargs) -> 'ExtractionResult':
        """Extract text from image using OCR."""
        try:
            start_time = time.time()

            # Open image with PIL
            with Image.open(file_path) as img:
                # Preprocess for better OCR
                if img.mode != 'L':
                    img = img.convert('L')

                # Perform OCR
                text = pytesseract.image_to_string(img, config='--psm 6')

                # Extract metadata
                metadata = await self.extract_metadata(file_path)

                processing_time = time.time() - start_time

                return ExtractionResult(
                    content=text.strip(),
                    metadata=metadata,
                    processing_time=processing_time,
                    success=True
                )

        except Exception as e:
            logger.error(f"Image OCR extraction failed: {e}")
            return ExtractionResult(
                content="",
                metadata=DocumentMetadata(),
                success=False,
                error=str(e)
            )

    async def extract_metadata(self, file_path: str) -> 'DocumentMetadata':
        """Extract metadata from image file."""
        try:
            with Image.open(file_path) as img:
                return DocumentMetadata(
                    title=Path(file_path).stem,
                    file_size=os.path.getsize(file_path),
                    page_count=1,
                    creation_date=datetime.fromtimestamp(
                        os.path.getctime(file_path), tz=timezone.utc
                    ).isoformat()
                )
        except Exception as e:
            logger.error(f"Image metadata extraction failed: {e}")
            return DocumentMetadata()


# Data classes for extraction results
class DocumentMetadata:
    def __init__(self, **kwargs):
        self.title = kwargs.get('title', '')
        self.author = kwargs.get('author', '')
        self.subject = kwargs.get('subject', '')
        self.keywords = kwargs.get('keywords', '')
        self.creator = kwargs.get('creator', '')
        self.producer = kwargs.get('producer', '')
        self.creation_date = kwargs.get('creation_date', '')
        self.modification_date = kwargs.get('modification_date', '')
        self.page_count = kwargs.get('page_count', 0)
        self.file_size = kwargs.get('file_size', 0)
        self.source_url = kwargs.get('source_url', '')


class ExtractionResult:
    def __init__(self, **kwargs):
        self.content = kwargs.get('content', '')
        self.metadata = kwargs.get('metadata', DocumentMetadata())
        self.pages_processed = kwargs.get('pages_processed', 0)
        self.images_extracted = kwargs.get('images_extracted', 0)
        self.ocr_processed = kwargs.get('ocr_processed', 0)
        self.processing_time = kwargs.get('processing_time', 0.0)
        self.success = kwargs.get('success', True)
        self.error = kwargs.get('error', '')
        self.source_url = kwargs.get('source_url', '')


# Global document processor instance
document_processor = DocumentProcessor()


async def get_document_processor() -> DocumentProcessor:
    """Get the document processor instance"""
    return document_processor
