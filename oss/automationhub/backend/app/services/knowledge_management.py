"""
Advanced Knowledge Management System for UPM.Plus

Provides comprehensive document processing, embedding generation,
intelligent chunking, and RAG-powered knowledge retrieval for the
conversational AI system.
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID, uuid4

import httpx
from pydantic import BaseModel, Field

from app.core.config import settings
from app.core.redis import redis_client
from app.services.vector_store import VectorStoreService
from app.services.llm import llm_service

logger = logging.getLogger(__name__)


class DocumentChunk(BaseModel):
    """Document chunk for vector storage."""
    id: str = Field(default_factory=lambda: str(uuid4()))
    document_id: UUID
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: Dict[str, Any] = Field(default_factory=dict)
    embedding: Optional[List[float]] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)


class ProcessedDocument(BaseModel):
    """Processed document with extracted content and metadata."""
    id: UUID = Field(default_factory=uuid4)
    filename: str
    content_type: str
    file_size: int
    original_content: str
    processed_content: str
    chunks: List[DocumentChunk] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)
    processing_status: str = "pending"  # pending, processing, completed, failed
    processing_error: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.utcnow)
    updated_at: datetime = Field(default_factory=datetime.utcnow)
    user_id: Optional[UUID] = None
    tags: List[str] = Field(default_factory=list)


class KnowledgeQuery(BaseModel):
    """Knowledge base query with context."""
    query: str
    context: Dict[str, Any] = Field(default_factory=dict)
    max_results: int = 10
    similarity_threshold: float = 0.7
    user_id: Optional[UUID] = None
    filters: Dict[str, Any] = Field(default_factory=dict)


class KnowledgeResult(BaseModel):
    """Knowledge query result."""
    query: str
    results: List[Dict[str, Any]]
    total_found: int
    processing_time_ms: int
    suggested_questions: List[str] = Field(default_factory=list)
    summary: Optional[str] = None


class KnowledgeManagementService:
    """
    Advanced knowledge management system with document processing,
    intelligent chunking, and RAG-powered retrieval.
    """

    def __init__(self):
        self.vector_store = VectorStoreService()
        self.documents: Dict[UUID, ProcessedDocument] = {}
        self.supported_types = {
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.json': 'application/json',
            '.csv': 'text/csv',
            '.html': 'text/html',
            '.py': 'text/x-python',
            '.js': 'text/javascript',
            '.ts': 'text/typescript'
        }

    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None,
        user_id: Optional[UUID] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> UUID:
        """Upload and process a document."""
        try:
            # Detect content type if not provided
            if not content_type:
                content_type, _ = mimetypes.guess_type(filename)
                if not content_type:
                    content_type = 'application/octet-stream'

            # Create document record
            document = ProcessedDocument(
                filename=filename,
                content_type=content_type,
                file_size=len(file_content),
                original_content="",  # Will be extracted
                processed_content="",
                user_id=user_id,
                metadata=metadata or {},
                processing_status="processing"
            )

            # Store document
            self.documents[document.id] = document

            # Start async processing
            asyncio.create_task(self._process_document_async(document, file_content))

            logger.info(f"Document upload initiated: {filename} ({document.id})")
            return document.id

        except Exception as e:
            logger.error(f"Failed to upload document: {e}")
            raise

    async def _process_document_async(self, document: ProcessedDocument, file_content: bytes):
        """Process document asynchronously."""
        try:
            # Extract text content
            text_content = await self._extract_text_content(file_content, document.content_type)
            document.original_content = text_content

            # Clean and preprocess content
            processed_content = await self._preprocess_content(text_content, document.metadata)
            document.processed_content = processed_content

            # Generate chunks
            chunks = await self._create_chunks(document, processed_content)
            document.chunks = chunks

            # Generate embeddings for chunks
            await self._generate_embeddings(chunks)

            # Store chunks in vector database
            await self._store_chunks_in_vector_db(chunks)

            # Extract metadata
            enhanced_metadata = await self._extract_metadata(document)
            document.metadata.update(enhanced_metadata)

            # Update status
            document.processing_status = "completed"
            document.updated_at = datetime.utcnow()

            # Cache in Redis
            await redis_client.set(
                f"document:{document.id}",
                json.dumps(document.dict(), default=str),
                expire=3600 * 24 * 7  # 7 days
            )

            logger.info(f"Document processing completed: {document.filename}")

        except Exception as e:
            document.processing_status = "failed"
            document.processing_error = str(e)
            document.updated_at = datetime.utcnow()
            logger.error(f"Document processing failed: {e}")

    async def _extract_text_content(self, file_content: bytes, content_type: str) -> str:
        """Extract text content from various file formats."""
        try:
            if content_type.startswith('text/'):
                return file_content.decode('utf-8', errors='ignore')

            elif content_type == 'application/pdf':
                return await self._extract_pdf_text(file_content)

            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._extract_docx_text(file_content)

            elif content_type == 'application/json':
                try:
                    data = json.loads(file_content.decode('utf-8'))
                    return json.dumps(data, indent=2)
                except:
                    return file_content.decode('utf-8', errors='ignore')

            elif content_type == 'text/csv':
                return file_content.decode('utf-8', errors='ignore')

            else:
                # Try to decode as text
                return file_content.decode('utf-8', errors='ignore')

        except Exception as e:
            logger.error(f"Failed to extract text content: {e}")
            return ""

    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF files."""
        try:
            import PyPDF2
            import io

            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""

            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return text

        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return ""

    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX files."""
        try:
            import docx
            import io

            doc = docx.Document(io.BytesIO(file_content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text

        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return ""

    async def _preprocess_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Clean and preprocess text content."""
        try:
            # Basic text cleaning
            lines = content.split('\n')
            processed_lines = []

            for line in lines:
                line = line.strip()
                if line and len(line) > 3:  # Filter out very short lines
                    processed_lines.append(line)

            processed_content = '\n'.join(processed_lines)

            # Remove excessive whitespace
            import re
            processed_content = re.sub(r'\n\s*\n', '\n\n', processed_content)
            processed_content = re.sub(r' +', ' ', processed_content)

            return processed_content

        except Exception as e:
            logger.error(f"Failed to preprocess content: {e}")
            return content

    async def _create_chunks(self, document: ProcessedDocument, content: str) -> List[DocumentChunk]:
        """Create intelligent chunks from document content."""
        try:
            chunks = []
            chunk_size = 1000  # Target chunk size in characters
            overlap_size = 200  # Overlap between chunks

            # Split by paragraphs first
            paragraphs = content.split('\n\n')
            current_chunk = ""
            chunk_index = 0
            start_char = 0

            for paragraph in paragraphs:
                # If adding this paragraph would exceed chunk size, create a chunk
                if len(current_chunk) + len(paragraph) > chunk_size and current_chunk:
                    chunk = DocumentChunk(
                        document_id=document.id,
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                        start_char=start_char,
                        end_char=start_char + len(current_chunk),
                        metadata={
                            "filename": document.filename,
                            "content_type": document.content_type,
                            "chunk_type": "paragraph_based",
                            "user_id": str(document.user_id) if document.user_id else None
                        }
                    )
                    chunks.append(chunk)

                    # Start new chunk with overlap
                    overlap_text = current_chunk[-overlap_size:] if len(current_chunk) > overlap_size else current_chunk
                    current_chunk = overlap_text + "\n\n" + paragraph
                    start_char = chunk.end_char - len(overlap_text)
                    chunk_index += 1
                else:
                    # Add paragraph to current chunk
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph

            # Add final chunk if any content remains
            if current_chunk.strip():
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=current_chunk.strip(),
                    chunk_index=chunk_index,
                    start_char=start_char,
                    end_char=start_char + len(current_chunk),
                    metadata={
                        "filename": document.filename,
                        "content_type": document.content_type,
                        "chunk_type": "paragraph_based",
                        "user_id": str(document.user_id) if document.user_id else None
                    }
                )
                chunks.append(chunk)

            logger.info(f"Created {len(chunks)} chunks for document {document.filename}")
            return chunks

        except Exception as e:
            logger.error(f"Failed to create chunks: {e}")
            return []

    async def _generate_embeddings(self, chunks: List[DocumentChunk]):
        """Generate embeddings for document chunks."""
        try:
            # For now, use a simple approach with sentence transformers
            # In production, would use OpenAI embeddings or similar
            for chunk in chunks:
                # Generate a simple hash-based embedding for testing
                # Replace with actual embedding generation
                content_hash = hashlib.md5(chunk.content.encode()).hexdigest()
                embedding = [float(int(content_hash[i:i+2], 16)) / 255.0 for i in range(0, 32, 2)]
                chunk.embedding = embedding

            logger.info(f"Generated embeddings for {len(chunks)} chunks")

        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")

    async def _store_chunks_in_vector_db(self, chunks: List[DocumentChunk]):
        """Store chunks in vector database."""
        try:
            if not chunks:
                return

            documents = [chunk.content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            ids = [chunk.id for chunk in chunks]
            embeddings = [chunk.embedding for chunk in chunks if chunk.embedding]

            success = await self.vector_store.add_documents(
                documents=documents,
                metadatas=metadatas,
                ids=ids,
                embeddings=embeddings if embeddings else None
            )

            if success:
                logger.info(f"Stored {len(chunks)} chunks in vector database")
            else:
                logger.error("Failed to store chunks in vector database")

        except Exception as e:
            logger.error(f"Failed to store chunks in vector database: {e}")

    async def _extract_metadata(self, document: ProcessedDocument) -> Dict[str, Any]:
        """Extract enhanced metadata from document."""
        try:
            metadata = {}

            # Basic statistics
            content = document.processed_content
            metadata["word_count"] = len(content.split())
            metadata["character_count"] = len(content)
            metadata["line_count"] = content.count('\n')

            # Use AI to extract key topics and summary
            if len(content) > 100:
                try:
                    analysis = await llm_service.analyze_task(
                        task_description=f"Analyze this document and extract key topics, entities, and a brief summary",
                        context=f"Document content (first 2000 chars): {content[:2000]}"
                    )

                    if isinstance(analysis, dict) and "analysis" in analysis:
                        analysis_result = analysis["analysis"]
                        if isinstance(analysis_result, dict):
                            metadata["ai_summary"] = analysis_result.get("summary", "")
                            metadata["key_topics"] = analysis_result.get("topics", [])
                            metadata["entities"] = analysis_result.get("entities", [])

                except Exception as e:
                    logger.warning(f"AI metadata extraction failed: {e}")

            # File-specific metadata
            if document.content_type == 'application/json':
                try:
                    data = json.loads(document.original_content)
                    metadata["json_structure"] = list(data.keys()) if isinstance(data, dict) else "array"
                except:
                    pass

            return metadata

        except Exception as e:
            logger.error(f"Failed to extract metadata: {e}")
            return {}

    async def query_knowledge(self, query: KnowledgeQuery) -> KnowledgeResult:
        """Query the knowledge base with semantic search."""
        start_time = datetime.utcnow()

        try:
            # Search vector database
            search_results = await self.vector_store.search_documents(
                query=query.query,
                n_results=query.max_results,
                where=query.filters
            )

            # Process results
            results = []
            for i, (doc, metadata, distance, doc_id) in enumerate(zip(
                search_results["documents"],
                search_results["metadatas"],
                search_results["distances"],
                search_results["ids"]
            )):
                if distance <= (1.0 - query.similarity_threshold):  # Convert similarity to distance
                    results.append({
                        "content": doc,
                        "metadata": metadata,
                        "similarity_score": 1.0 - distance,
                        "chunk_id": doc_id,
                        "rank": i + 1
                    })

            # Generate summary and suggested questions using AI
            summary = None
            suggested_questions = []

            if results:
                try:
                    # Create context from top results
                    context_text = "\n\n".join([r["content"] for r in results[:3]])

                    analysis = await llm_service.analyze_task(
                        task_description=f"Based on the search results, provide a summary and suggest 3 follow-up questions for: {query.query}",
                        context=f"Search results: {context_text[:1500]}"
                    )

                    if isinstance(analysis, dict) and "analysis" in analysis:
                        analysis_result = analysis["analysis"]
                        if isinstance(analysis_result, dict):
                            summary = analysis_result.get("summary", "")
                            suggested_questions = analysis_result.get("questions", [])

                except Exception as e:
                    logger.warning(f"AI analysis for query results failed: {e}")

            # Calculate processing time
            processing_time = int((datetime.utcnow() - start_time).total_seconds() * 1000)

            return KnowledgeResult(
                query=query.query,
                results=results,
                total_found=len(results),
                processing_time_ms=processing_time,
                suggested_questions=suggested_questions,
                summary=summary
            )

        except Exception as e:
            logger.error(f"Knowledge query failed: {e}")
            processing_time = int((datetime.utcnow() - start_time).total_seconds() * 1000)

            return KnowledgeResult(
                query=query.query,
                results=[],
                total_found=0,
                processing_time_ms=processing_time,
                summary=f"Query failed: {str(e)}"
            )

    async def get_document(self, document_id: UUID) -> Optional[ProcessedDocument]:
        """Get document by ID."""
        return self.documents.get(document_id)

    async def list_documents(
        self,
        user_id: Optional[UUID] = None,
        status: Optional[str] = None,
        limit: int = 100
    ) -> List[ProcessedDocument]:
        """List documents with optional filters."""
        documents = list(self.documents.values())

        if user_id:
            documents = [d for d in documents if d.user_id == user_id]

        if status:
            documents = [d for d in documents if d.processing_status == status]

        # Sort by creation date
        documents.sort(key=lambda x: x.created_at, reverse=True)

        return documents[:limit]

    async def delete_document(self, document_id: UUID) -> bool:
        """Delete document and its chunks."""
        try:
            document = self.documents.get(document_id)
            if not document:
                return False

            # Delete chunks from vector database
            if document.chunks:
                chunk_ids = [chunk.id for chunk in document.chunks]
                await self.vector_store.delete_documents(chunk_ids)

            # Remove from memory
            del self.documents[document_id]

            # Remove from Redis cache
            await redis_client.delete(f"document:{document_id}")

            logger.info(f"Deleted document: {document.filename}")
            return True

        except Exception as e:
            logger.error(f"Failed to delete document: {e}")
            return False

    async def update_document_metadata(
        self,
        document_id: UUID,
        metadata: Dict[str, Any],
        tags: Optional[List[str]] = None
    ) -> bool:
        """Update document metadata and tags."""
        try:
            document = self.documents.get(document_id)
            if not document:
                return False

            # Update metadata
            document.metadata.update(metadata)

            # Update tags
            if tags is not None:
                document.tags = tags

            document.updated_at = datetime.utcnow()

            # Update cache
            await redis_client.set(
                f"document:{document_id}",
                json.dumps(document.dict(), default=str),
                expire=3600 * 24 * 7
            )

            logger.info(f"Updated document metadata: {document.filename}")
            return True

        except Exception as e:
            logger.error(f"Failed to update document metadata: {e}")
            return False

    async def get_statistics(self, user_id: Optional[UUID] = None) -> Dict[str, Any]:
        """Get knowledge base statistics."""
        try:
            documents = list(self.documents.values())

            if user_id:
                documents = [d for d in documents if d.user_id == user_id]

            # Basic statistics
            total_documents = len(documents)
            completed_documents = len([d for d in documents if d.processing_status == "completed"])
            total_chunks = sum(len(d.chunks) for d in documents)
            total_size = sum(d.file_size for d in documents)

            # Content type distribution
            content_types = {}
            for doc in documents:
                content_types[doc.content_type] = content_types.get(doc.content_type, 0) + 1

            # Vector database info
            vector_info = await self.vector_store.get_collection_info()

            return {
                "documents": {
                    "total": total_documents,
                    "completed": completed_documents,
                    "processing": len([d for d in documents if d.processing_status == "processing"]),
                    "failed": len([d for d in documents if d.processing_status == "failed"])
                },
                "chunks": {
                    "total": total_chunks,
                    "in_vector_db": vector_info.get("count", 0)
                },
                "storage": {
                    "total_size_bytes": total_size,
                    "average_size_bytes": total_size // total_documents if total_documents > 0 else 0
                },
                "content_types": content_types,
                "vector_database": vector_info
            }

        except Exception as e:
            logger.error(f"Failed to get statistics: {e}")
            return {"error": str(e)}


# Global knowledge management service instance
knowledge_service = KnowledgeManagementService()